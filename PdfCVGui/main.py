import os
import sys
import itertools
from pathlib import Path

import docx
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT

from PdfCVGui.ocr import TesseractOCR, PDF

if "MEI" not in str(Path(__file__).parent):
    PDFS = Path(__file__).parent.parent / "pdfs"
    DATA = Path(__file__).parent.parent.parent / "data"
    RESULTS = Path(__file__).parent.parent / "results"
else:
    PDFS = Path(sys.executable).parent / "pdfs"
    RESULTS = Path(sys.executable).parent / "results"
    TEMP = Path(sys.executable).parent / "temp"

def run_main(filepath):
    if not os.path.exists(RESULTS):
        os.mkdir(RESULTS)
    convert_pdf_to_docx(filepath, os.path.basename(filepath))


def convert_pdf_to_docx(path, name):
    ocr = TesseractOCR(3, "eng")
    pdf = PDF(str(path))
    tables = pdf.extract_tables(
        ocr=ocr,
        implicit_rows=True,
        borderless_tables=True
    )
    doc = docx.Document()
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_height, section.page_width = section.page_width + 10, section.page_height + 25
    for idx, tables in tables.items():
        for table in tables:
            add_table(doc, table)
    doc.save(str(RESULTS / (name + ".docx")))

def add_table(doc, table):
    data = {}
    maxr = maxc = 0
    for r, row in table.content.items():
        if r > maxr:
            maxr = r
        for c, col in enumerate(row):
            if c > maxc:
                maxc = c
            data.setdefault(hash(col), [])
            data[hash(col)] += [(col, r, c)]
    tab = doc.add_table(cols=maxc+1, rows=maxr+ 1)
    tab.style = "Table Grid"
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    tab.autofit = False
    tab.allow_autofit = False
    for row in tab.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Pt(0)
    for c in data.values():
        if len(c) == 1:
            pos = c.pop()
            if pos[0].value is not None:
                tab.cell(pos[1], pos[2]).text = pos[0].value
        else:
            cell1 = None
            for c1, c2 in itertools.pairwise(c):
                if [c2[1], c2[2]] in gen_neighbors(c1):
                    row1 = tab.rows[c1[1]]
                    row2 = tab.rows[c2[1]]
                    cell1 = row1.cells[c1[2]]
                    cell2 = row2.cells[c2[2]]
                    try:
                        cell1.merge(cell2)
                    except:
                        pass
            if c1[0].value is not None and cell1 is not None:
                cell1.text = c1[0].value
    p = doc.add_paragraph()
    p.add_run("")

def gen_neighbors(c1):
    for x,y in [(1,0), (0,1), (-1,0), (-1,1), (1,-1), (0,-1), (-1,-1), (1,1)]:
        yield [c1[1] + x, c1[2] + y]
